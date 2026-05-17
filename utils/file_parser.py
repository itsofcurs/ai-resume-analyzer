"""
File Parser Utility
Extracts raw text from PDF, DOCX, and TXT files.
Supports: PyMuPDF (fitz) for PDF, python-docx for DOCX.
"""

import os
import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    Falls back to pdfplumber if fitz is unavailable.
    """
    try:
        import fitz  # PyMuPDF
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("PyMuPDF not found, trying pdfplumber...")
    except Exception as e:
        logger.error(f"PDF extraction error (fitz): {e}")

    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber not found either.")
    except Exception as e:
        logger.error(f"PDF extraction error (pdfplumber): {e}")

    return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Preserves paragraph structure.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extract text from a TXT file.
    Tries UTF-8, then latin-1 encoding.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Main entry point: detect file type and extract text accordingly.
    Returns raw text string.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        text = extract_text_from_txt(file_bytes)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        text = ""

    # Basic clean-up: remove excessive whitespace
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)
